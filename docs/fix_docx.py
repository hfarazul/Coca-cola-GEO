"""Post-process pandoc-generated docx: add table borders, fix spacing."""
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_LINE_SPACING


def set_table_borders(table):
    """Add solid black borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    # Remove existing borders if any
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def shade_header_row(table):
    """Apply light gray shading to the first row (header)."""
    for cell in table.rows[0].cells:
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True


def style_table_text(table):
    """Set table cell text to 9pt with tight spacing."""
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    run.font.size = Pt(9)


def fix_paragraph_spacing(doc):
    """Add breathing room between paragraphs."""
    for para in doc.paragraphs:
        fmt = para.paragraph_format
        style_name = para.style.name if para.style else ""

        if style_name.startswith("Heading 1"):
            fmt.space_before = Pt(24)
            fmt.space_after = Pt(8)
            fmt.line_spacing = 1.15
            for run in para.runs:
                run.font.size = Pt(22)
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        elif style_name.startswith("Heading 2"):
            fmt.space_before = Pt(18)
            fmt.space_after = Pt(6)
            fmt.line_spacing = 1.15
            for run in para.runs:
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0x2B, 0x2B, 0x4E)
        elif style_name.startswith("Heading 3"):
            fmt.space_before = Pt(14)
            fmt.space_after = Pt(4)
            fmt.line_spacing = 1.15
            for run in para.runs:
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(0x3C, 0x3C, 0x6E)
        else:
            # Body text
            fmt.space_before = Pt(3)
            fmt.space_after = Pt(6)
            fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            fmt.line_spacing = 1.25
            for run in para.runs:
                if run.font.size is None:
                    run.font.size = Pt(10.5)


def set_margins(doc):
    """Set page margins for a less cramped look."""
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


def main(input_path, output_path):
    doc = Document(input_path)
    set_margins(doc)
    fix_paragraph_spacing(doc)

    for table in doc.tables:
        set_table_borders(table)
        shade_header_row(table)
        style_table_text(table)

    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    inp = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else inp
    main(inp, out)
