#!/usr/bin/env python3
"""Convert .md to Google Docs-ready .docx with proper table borders and spacing.

Usage:
    python md2docx.py input.md                  # outputs input.docx
    python md2docx.py input.md output.docx      # custom output path
"""
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH


def convert_md_to_docx(md_path: str, docx_path: str):
    """Full pipeline: md → pandoc → post-process → final docx."""
    md_dir = str(Path(md_path).parent)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name

    # Step 1: Pandoc conversion (run from md directory so image paths resolve)
    subprocess.run(
        ["pandoc", md_path, "-o", tmp_path, "--from=markdown", "--to=docx",
         f"--resource-path={md_dir}"],
        check=True,
    )

    # Step 2: Post-process with python-docx
    doc = Document(tmp_path)
    _set_margins(doc)
    _fix_paragraph_spacing(doc)
    for table in doc.tables:
        _set_table_full_width(table)
        _shade_header_row(table)
        _style_table_text(table)

    # Step 3: Add logo to header and footer of every page
    logo_path = Path(md_dir) / "cortivo-logo.png"
    if logo_path.exists():
        _add_header_logo(doc, str(logo_path))
        _add_footer_logo(doc, str(logo_path))

    doc.save(docx_path)

    Path(tmp_path).unlink(missing_ok=True)
    print(f"Created: {docx_path}")


# --- Styling functions ---


def _add_header_logo(doc, logo_path: str):
    """Add Cortivo logo + 'Cortivo AI' text, top-left, vertically aligned using a borderless table."""
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    first_para = doc.paragraphs[0]

    # Create a 1-row, 2-col table before the first paragraph
    tbl_xml = (
        f'<w:tbl {nsdecls("w")}>'
        '<w:tblPr>'
        f'  <w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'
        f'  <w:tblBorders {nsdecls("w")}>'
        '    <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '    <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '    <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '    <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '    <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '    <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  </w:tblBorders>'
        '</w:tblPr>'
        '<w:tblGrid><w:gridCol w:w="800"/><w:gridCol w:w="8560"/></w:tblGrid>'
        '<w:tr>'
        '  <w:tc><w:tcPr><w:tcW w:w="800" w:type="dxa"/><w:vAlign w:val="center"/></w:tcPr>'
        '    <w:p><w:pPr><w:spacing w:after="0" w:before="0"/></w:pPr></w:p>'
        '  </w:tc>'
        '  <w:tc><w:tcPr><w:tcW w:w="8560" w:type="dxa"/><w:vAlign w:val="center"/></w:tcPr>'
        '    <w:p><w:pPr><w:spacing w:after="0" w:before="0"/></w:pPr></w:p>'
        '  </w:tc>'
        '</w:tr>'
        '</w:tbl>'
    )
    tbl_element = parse_xml(tbl_xml)
    first_para._element.addprevious(tbl_element)

    # Also add a spacer paragraph between the table and the title
    spacer = parse_xml(
        f'<w:p {nsdecls("w")}><w:pPr>'
        '<w:spacing w:after="120" w:before="0"/></w:pPr></w:p>'
    )
    first_para._element.addprevious(spacer)

    # Now populate the table cells via the docx Table object
    from docx.table import Table as DocxTable
    table = DocxTable(tbl_element, doc)

    # Left cell: logo
    logo_para = table.rows[0].cells[0].paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_run = logo_para.add_run()
    logo_run.add_picture(logo_path, height=Inches(0.4))

    # Right cell: "Cortivo AI" text
    text_para = table.rows[0].cells[1].paragraphs[0]
    text_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    text_run = text_para.add_run("Cortivo AI")
    text_run.font.size = Pt(18)
    text_run.font.bold = True
    text_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def _add_footer_logo(doc, logo_path: str):
    """Add the Cortivo logo centered in the footer of every page."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        # Clear any existing footer content
        for p in footer.paragraphs:
            p.clear()
        # Use first paragraph or create one
        if footer.paragraphs:
            para = footer.paragraphs[0]
        else:
            para = footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run()
        run.add_picture(logo_path, height=Inches(0.35))


def _set_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


def _fix_paragraph_spacing(doc):
    for para in doc.paragraphs:
        fmt = para.paragraph_format
        style_name = para.style.name if para.style else ""

        if style_name.startswith("Heading 1"):
            fmt.space_before = Pt(28)
            fmt.space_after = Pt(12)
            fmt.line_spacing = 1.15
            for run in para.runs:
                run.font.size = Pt(24)
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        elif style_name.startswith("Heading 2"):
            fmt.space_before = Pt(24)
            fmt.space_after = Pt(10)
            fmt.line_spacing = 1.15
            for run in para.runs:
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(0x2B, 0x2B, 0x4E)

        elif style_name.startswith("Heading 3"):
            fmt.space_before = Pt(18)
            fmt.space_after = Pt(8)
            fmt.line_spacing = 1.15
            for run in para.runs:
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0x3C, 0x3C, 0x6E)

        else:
            fmt.space_before = Pt(3)
            fmt.space_after = Pt(6)
            fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            fmt.line_spacing = 1.3
            for run in para.runs:
                if run.font.size is None:
                    run.font.size = Pt(12)


def _guess_col_widths(table, page_width=9360):
    """Calculate proportional column widths based on header content length."""
    num_cols = len(table.columns)
    if num_cols == 0:
        return []

    # Sample max content length per column from all rows
    max_lens = [0] * num_cols
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            text = cell.text.strip()
            max_lens[i] = max(max_lens[i], len(text))

    # Clamp: text columns (long) get more, short % columns get a minimum
    MIN_COL = 600   # ~0.4 inch minimum for narrow cols
    weights = []
    for length in max_lens:
        # Square root scaling so long text cols get more but don't dominate
        weights.append(max(length ** 0.7, 2))

    total_weight = sum(weights)
    widths = [max(int(page_width * w / total_weight), MIN_COL) for w in weights]

    # Adjust to hit exact page_width
    diff = page_width - sum(widths)
    # Give remainder to the widest column
    widest_idx = widths.index(max(widths))
    widths[widest_idx] += diff

    return widths


def _set_table_full_width(table):
    """Force table to span full page width with proportional columns."""
    tbl = table._tbl
    page_width = 9360  # twips (6.5 inches)
    num_cols = len(table.columns)
    if num_cols == 0:
        return

    col_widths = _guess_col_widths(table, page_width)

    # --- Rewrite tblGrid ---
    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    grid_xml = f'<w:tblGrid {nsdecls("w")}>'
    for w in col_widths:
        grid_xml += f'<w:gridCol w:w="{w}"/>'
    grid_xml += "</w:tblGrid>"
    new_grid = parse_xml(grid_xml)
    tblPr = tbl.tblPr
    if tblPr is not None:
        tblPr.addnext(new_grid)
    else:
        tbl.insert(0, new_grid)

    # --- tblPr ---
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)

    for tag in ("w:tblW", "w:tblLayout", "w:tblBorders", "w:tblCellMar", "w:tblInd"):
        el = tblPr.find(qn(tag))
        if el is not None:
            tblPr.remove(el)

    tblPr.append(parse_xml(
        f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'
    ))
    tblPr.append(parse_xml(
        f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>'
    ))
    tblPr.append(parse_xml(
        f'<w:tblInd {nsdecls("w")} w:w="0" w:type="dxa"/>'
    ))
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="6" w:space="0" w:color="4472C4"/>'
        '  <w:left w:val="single" w:sz="6" w:space="0" w:color="4472C4"/>'
        '  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="4472C4"/>'
        '  <w:right w:val="single" w:sz="6" w:space="0" w:color="4472C4"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="8EAADB"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="8EAADB"/>'
        "</w:tblBorders>"
    ))
    tblPr.append(parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>'
        '  <w:top w:w="60" w:type="dxa"/>'
        '  <w:left w:w="100" w:type="dxa"/>'
        '  <w:bottom w:w="60" w:type="dxa"/>'
        '  <w:right w:w="100" w:type="dxa"/>'
        "</w:tblCellMar>"
    ))

    # --- Set each cell width to match its column ---
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is not None:
                tcPr.remove(tcW)
            tcPr.append(parse_xml(
                f'<w:tcW {nsdecls("w")} w:w="{col_widths[i]}" w:type="dxa"/>'
            ))


def _shade_header_row(table):
    if not table.rows:
        return
    for cell in table.rows[0].cells:
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="4472C4" w:val="clear"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(11)


def _style_table_text(table):
    for i, row in enumerate(table.rows):
        if i == 0:
            continue  # Header already styled
        # Alternate row shading
        for cell in row.cells:
            if i % 2 == 0:
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="D6E4F0" w:val="clear"/>'
                )
                cell._tc.get_or_add_tcPr().append(shading)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.1
                for run in p.runs:
                    run.font.size = Pt(11)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python md2docx.py input.md [output.docx]")
        sys.exit(1)

    md_file = sys.argv[1]
    if len(sys.argv) > 2:
        docx_file = sys.argv[2]
    else:
        docx_file = str(Path(md_file).with_suffix(".docx"))

    convert_md_to_docx(md_file, docx_file)
